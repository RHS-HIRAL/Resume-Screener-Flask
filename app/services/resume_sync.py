"""
app/services/resume_sync.py
Resume sync pipelines:
  Pipeline 1 — Fetch emails from Outlook → download resumes → upload to SharePoint.
               Stamps Source="Website" on every uploaded resume.
  Pipeline 2 — Extract text from all SharePoint resumes → upload .txt files.
              Corrupted files (no extractable text) get MatchScore = -1.
"""

import io
import os
import re
import shutil
from datetime import datetime, timedelta, timezone
from pathlib import Path

import requests
from bs4 import BeautifulSoup

from config import Config
from app.services.sharepoint_sync import (
    GraphAuthProvider,
    SyncSharePointManager,
    CandidateInfo,
    SUBJECT_PATTERN,
    JOB_ID_BRACKET,
    FIELD_PATTERNS,
    _unique_base_path,
    logger,
)

# ── Optional dependencies ──────────────────────────────────────────────────────
try:
    import fitz as pymupdf

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pytesseract
    from PIL import Image

    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import PyPDF2

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ══════════════════════════════════════════════════════════════════════════════
#  EMAIL FETCHER
# ══════════════════════════════════════════════════════════════════════════════


class EmailFetcher:
    def __init__(self, headers: dict):
        self.headers = headers
        self.base = "https://graph.microsoft.com/v1.0"
        # OPTIMIZATION: Implement requests Session to reuse TCP connections
        self.session = requests.Session()
        self.session.headers.update(headers)

    def fetch_recent_emails(self) -> list:
        since = (
            datetime.now(timezone.utc) - timedelta(hours=Config.MAILBOX_LOOKBACK_HOURS)
        ).strftime("%Y-%m-%dT%H:%M:%SZ")
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}/messages"
            f"?$filter=receivedDateTime ge {since}"
            f"&$orderby=receivedDateTime desc&$top=200"
            f"&$select=id,subject,from,receivedDateTime,body,hasAttachments"
        )
        all_msgs = []
        while url:
            resp = self.session.get(url, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            all_msgs.extend(data.get("value", []))
            url = data.get("@odata.nextLink")

        keywords = Config.RESUME_SUBJECT_KEYWORDS
        relevant = [
            m
            for m in all_msgs
            if any(kw in (m.get("subject") or "").lower() for kw in keywords)
        ]
        logger.info(
            "Fetched %d relevant emails (of %d total).", len(relevant), len(all_msgs)
        )
        return [self._parse_email(m) for m in relevant]

    def get_attachment_content(self, email_id: str, attachment_id: str) -> bytes:
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}"
            f"/messages/{email_id}/attachments/{attachment_id}/$value"
        )
        resp = self.session.get(url, timeout=60)
        resp.raise_for_status()
        return resp.content

    def get_attachments_metadata(self, email_id: str) -> list:
        url = (
            f"{self.base}/users/{Config.MAILBOX_USER}"
            f"/messages/{email_id}/attachments?$select=id,name,contentType"
        )
        resp = self.session.get(url, timeout=30)
        resp.raise_for_status()
        return resp.json().get("value", [])

    def _parse_email(self, msg: dict) -> CandidateInfo:
        subject = msg.get("subject", "")
        body_text = BeautifulSoup(
            msg.get("body", {}).get("content", ""), "html.parser"
        ).get_text(separator="\n")
        received_dt = msg.get("receivedDateTime", "")

        candidate = CandidateInfo(
            source_email_id=msg["id"],
            source_subject=subject,
            received_datetime=received_dt,
        )

        subj_match = SUBJECT_PATTERN.search(subject)
        if subj_match:
            candidate.job_role = subj_match.group(1).strip()
            candidate.job_id = subj_match.group(2).strip()

        for line in body_text.splitlines():
            line = line.strip()
            if not line:
                continue
            for fk, pattern in FIELD_PATTERNS.items():
                m = pattern.search(line)
                if not m:
                    continue
                value = m.group(1).strip()
                if fk == "name":
                    candidate.name = value.title()
                elif fk == "email":
                    candidate.email = value.lower()
                elif fk == "phone":
                    candidate.phone = re.sub(r"[^\d+\-() ]", "", value).strip()
                elif fk == "resume_url":
                    candidate.resume_url = value
                elif fk == "job_opening" and not candidate.job_id:
                    bracket = JOB_ID_BRACKET.search(value)
                    if bracket:
                        candidate.job_id = bracket.group(1)
                        candidate.job_role = value[: bracket.start()].strip()
                    else:
                        candidate.job_role = candidate.job_role or value

        if not candidate.resume_url and msg.get("hasAttachments"):
            atts = self.get_attachments_metadata(msg["id"])
            valid_ext = (".pdf", ".docx", ".doc")
            valid_mime = ["pdf", "wordprocessingml", "msword"]
            candidate.attachments = [
                {
                    "id": a["id"],
                    "name": a["name"],
                    "content_type": a.get("contentType", ""),
                }
                for a in atts
                if a.get("name", "").lower().endswith(valid_ext)
                or any(ct in a.get("contentType", "").lower() for ct in valid_mime)
            ]

        if not candidate.name:
            sender = msg.get("from", {}).get("emailAddress", {}).get("name", "")
            candidate.name = sender.title() if sender else "Unknown"
        if not candidate.email:
            candidate.email = (
                msg.get("from", {}).get("emailAddress", {}).get("address", "")
            )

        return candidate


# ══════════════════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════


def extract_text_from_pdf(file_path: str) -> str:
    if not HAS_PYPDF2:
        return ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.error("PyPDF2 failed for %s: %s", file_path, e)
        return ""


def extract_text_from_docx(file_path: str) -> str:
    if not HAS_DOCX:
        return ""
    try:
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        logger.error("DOCX extraction failed for %s: %s", file_path, e)
        return ""


def extract_text_with_ocr(pdf_path: str) -> str:
    if not HAS_PYMUPDF or not HAS_TESSERACT:
        return ""
    if Config.TESSERACT_CMD and os.path.exists(Config.TESSERACT_CMD):
        pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_CMD
    try:
        doc = pymupdf.open(pdf_path)
        pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=Config.OCR_DPI)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            pages.append(pytesseract.image_to_string(img))
        doc.close()
        return "\n\n".join(pages)
    except Exception as e:
        logger.error("OCR failed: %s", e)
        return ""


def extract_raw_text(local_path: Path) -> str:
    suffix = local_path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text_from_pdf(str(local_path))
    elif suffix in (".docx", ".doc"):
        text = extract_text_from_docx(str(local_path))
    else:
        return ""

    if suffix == ".pdf" and (not text or len(text.strip()) < 10):
        ocr = extract_text_with_ocr(str(local_path))
        if ocr and len(ocr.strip()) >= 10:
            return ocr
    return text


# ══════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════════════════


def _download_resume_from_url(url: str, dest_base: str):
    try:
        resp = requests.get(url, timeout=60, stream=True)
        resp.raise_for_status()
        ct = resp.headers.get("Content-Type", "").lower()
        ext = ".pdf"
        if "wordprocessingml" in ct or "msword" in ct:
            ext = ".docx"
        elif ".docx" in url.lower():
            ext = ".docx"

        tmp = dest_base + ext
        with open(tmp, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)

        with open(tmp, "rb") as f:
            header = f.read(4)
        if header.startswith(b"PK"):
            ext = ".docx"

        final = dest_base + ext
        if tmp != final:
            os.rename(tmp, final)
        return True, final, ext
    except Exception as e:
        logger.error("Failed to download from %s: %s", url, e)
        return False, "", ""


# ══════════════════════════════════════════════════════════════════════════════
#  PIPELINE 1 — Email Fetch & Resume Upload
# ══════════════════════════════════════════════════════════════════════════════


def run_email_fetch_pipeline(auth: GraphAuthProvider) -> dict:
    logger.info("=== PIPELINE 1: EMAIL FETCH & RESUME UPLOAD ===")
    headers = auth.get_headers()
    fetcher = EmailFetcher(headers)
    candidates = fetcher.fetch_recent_emails()

    if not candidates:
        logger.info("No new application emails found.")
        return {"success": 0, "failed": 0, "skipped": 0}

    seen = set()
    unique = []
    for c in candidates:
        key = (c.email.lower(), c.job_id)
        if key not in seen:
            seen.add(key)
            unique.append(c)

    os.makedirs(Config.SYNC_TEMP_RESUMES_DIR, exist_ok=True)
    sp = SyncSharePointManager(headers)
    base_folder = Config.SHAREPOINT_JOBS_FOLDER.strip("/")
    drive_id = sp._get_drive_id()
    sp._ensure_folder(drive_id, base_folder)

    results = {"success": 0, "failed": 0, "skipped": 0}
    candidates_by_folder = {}

    for c in unique:
        subfolder = f"{c.safe_job_id}_{c.safe_job_role}"
        candidates_by_folder.setdefault(subfolder, []).append(c)

    for subfolder, folder_candidates in candidates_by_folder.items():
        full_sp_folder = f"{base_folder}/{subfolder}"

        try:
            existing_folder_files = sp.get_folder_files_metadata(full_sp_folder)
        except AttributeError:
            existing_folder_files = None

        for candidate in folder_candidates:
            local_base = _unique_base_path(Config.SYNC_TEMP_RESUMES_DIR, candidate)
            exts_to_check = [".pdf", ".docx", ".doc"]

            if candidate.attachments:
                _, att_ext = os.path.splitext(candidate.attachments[0]["name"])
                if att_ext:
                    exts_to_check = [att_ext.lower()]

            already_processed = False
            for check_ext in exts_to_check:
                temp_filename = (
                    f"{candidate.safe_name}_{candidate.safe_job_id}{check_ext}"
                )
                if existing_folder_files is not None:
                    if (
                        temp_filename in existing_folder_files
                        and existing_folder_files[temp_filename]
                        == candidate.source_email_id
                    ):
                        already_processed = True
                        break
                else:
                    existing = sp.get_file_metadata(full_sp_folder, temp_filename)
                    if (
                        existing
                        and existing.get("SourceEmailID") == candidate.source_email_id
                    ):
                        already_processed = True
                        break

            if already_processed:
                logger.info(
                    "Already processed email for %s. Skipping without downloading.",
                    candidate.name,
                )
                results["skipped"] += 1
                continue

            downloaded, local_path, ext = False, "", ".pdf"

            if candidate.resume_url:
                downloaded, local_path, ext = _download_resume_from_url(
                    candidate.resume_url, local_base
                )

            if not downloaded and candidate.attachments:
                att = candidate.attachments[0]
                _, att_ext = os.path.splitext(att["name"])
                ext = att_ext.lower() or ".pdf"
                local_path = local_base + ext
                try:
                    content = fetcher.get_attachment_content(
                        candidate.source_email_id, att["id"]
                    )
                    with open(local_path, "wb") as f:
                        f.write(content)
                    downloaded = True
                except Exception as e:
                    logger.error("Attachment download failed: %s", e)

            if not downloaded:
                results["skipped"] += 1
                continue

            target_filename = f"{candidate.safe_name}_{candidate.safe_job_id}{ext}"

            # CHANGED: Removed JobID and JobRole (resume is already in the correct job-role subfolder). Added Source="Website" to identify Outlook-sourced resumes and prevent them from being renamed after screening.
            metadata = {
                "CandidateName": candidate.name,
                "CandidateEmail": candidate.email,
                "CandidatePhone": candidate.phone,
                "SourceEmailID": candidate.source_email_id,
                "Source": "Website",
            }

            try:
                sp.upload_resume(local_path, target_filename, subfolder, metadata)
                results["success"] += 1
                if existing_folder_files is not None:
                    existing_folder_files[target_filename] = candidate.source_email_id
            except Exception as e:
                results["failed"] += 1
                logger.error("Upload failed for %s: %s", candidate.name, e)
            finally:
                if local_path and os.path.exists(local_path):
                    try:
                        os.remove(local_path)
                    except OSError:
                        pass

    logger.info(
        "Pipeline 1 done. Uploaded: %d | Skipped: %d | Failed: %d",
        results["success"],
        results["skipped"],
        results["failed"],
    )
    return results


# ══════════════════════════════════════════════════════════════════════════════
#  PIPELINE 2 — Text Extraction & Upload (with corrupted file handling)
# ══════════════════════════════════════════════════════════════════════════════


def run_text_extraction_pipeline(auth: GraphAuthProvider) -> dict:
    logger.info("=== PIPELINE 2: TEXT EXTRACTION ===")
    headers = auth.get_headers()
    sp = SyncSharePointManager(headers)
    tmp_dir = Path(Config.SYNC_TEMP_RESUMES_DIR) / "text_extraction"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    base_folder = Config.SHAREPOINT_JOBS_FOLDER.strip("/")
    folders = sp.list_subfolders(base_folder)

    total_uploaded, total_skipped, total_failed = 0, 0, 0

    for folder in sorted(folders, key=lambda f: f["name"]):
        folder_path = f"{base_folder}/{folder['name']}"

        # OPTIMIZATION: Retrieve all files once to avoid O(N) calls for .txt existence
        all_items = sp.list_files(
            folder_path
        )  # Pass without extensions to grab everything
        if not all_items:
            continue

        resumes = []
        existing_txts = set()

        for item in all_items:
            fname_lower = item["name"].lower()
            if fname_lower.endswith((".pdf", ".docx", ".doc")):
                resumes.append(item)
            elif fname_lower.endswith(".txt"):
                existing_txts.add(fname_lower)

        for res in resumes:
            fname = res["name"]
            txt_filename = Path(fname).stem + ".txt"

            # O(1) Memory check instead of HTTP network lookup
            if txt_filename.lower() in existing_txts:
                total_skipped += 1
                continue

            local_resume_path = tmp_dir / fname
            local_txt_path = tmp_dir / txt_filename

            try:
                dl_url = res.get("download_url", "")
                ok = (
                    sp.download_file_by_url(dl_url, local_resume_path)
                    if dl_url
                    else sp.download_file(res["id"], local_resume_path)
                )

                if not ok:
                    total_failed += 1
                    continue

                text = extract_raw_text(local_resume_path)

                if not text or len(text.strip()) < 10:
                    logger.warning(
                        "Corrupted file detected: %s. Setting MatchScore=-1.", fname
                    )
                    # OPTIMIZATION: Re-use cached item ID instead of Graph API search
                    sp.update_match_score(res["id"], -1)
                    total_failed += 1
                else:
                    local_txt_path.write_text(text, encoding="utf-8")
                    if sp.upload_text_file(
                        local_txt_path, f"{folder_path}/{txt_filename}"
                    ):
                        total_uploaded += 1
                    else:
                        total_skipped += 1

            except Exception as e:
                logger.error("Failed for %s: %s", fname, e)
                total_failed += 1
            finally:
                # OPTIMIZATION: Ensure rigid disk cleanup regardless of exceptions
                local_resume_path.unlink(missing_ok=True)
                local_txt_path.unlink(missing_ok=True)

    try:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except OSError:
        pass

    logger.info(
        "Pipeline 2 done. Uploaded: %d | Skipped: %d | Failed: %d",
        total_uploaded,
        total_skipped,
        total_failed,
    )
    return {
        "uploaded": total_uploaded,
        "skipped": total_skipped,
        "failed": total_failed,
    }
